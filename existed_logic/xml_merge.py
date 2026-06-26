import time
import os
from docx import Document


def merge_docx_advanced(files, output_file, method='composer_restart'):
    """
    Merge DOCX files with format preservation

    Methods:
    - 'composer_restart': Try docxcompose, restart master on error (RECOMMENDED - Most reliable)
    - 'composer_safe': Use docxcompose with safe section handling (Fast)
    - 'paragraph_api': Use python-docx API (Safest, no corruption)

    Priority: No corruption > Format preservation > Speed
    """
    total_start_time = time.time()

    print(f"\n{'='*60}")
    print(f"BẮT ĐẦU GHÉP {len(files)} FILE - Phương thức: {method}")
    print(f"{'='*60}\n")

    if method == 'composer_restart':
        return _merge_composer_restart(files, output_file, total_start_time)
    elif method == 'paragraph_api':
        return _merge_paragraph_api(files, output_file, total_start_time)
    else:
        return _merge_composer_safe(files, output_file, total_start_time)


def _merge_composer_restart(files, output_file, total_start_time):
    """
    Method 1: Use docxcompose with restart on error
    PROS: Most reliable, preserves all formatting, handles relationships properly
    CONS: Slightly slower due to potential restarts
    """
    try:
        from docxcompose.composer import Composer

        merged_files = []

        # 1. Load master
        t_start = time.time()
        print(f"[1/{len(files)}] Đang tải file gốc: {os.path.basename(files[0])}...", end=" ", flush=True)

        master = Document(files[0])
        composer = Composer(master)
        merged_files.append(files[0])

        print(f"✓ ({time.time() - t_start:.2f}s)")

        # 2. Try to merge each file
        for i, file_path in enumerate(files[1:], 2):
            t_file_start = time.time()
            print(f"[{i}/{len(files)}] Đang ghép: {os.path.basename(file_path)}...", end=" ", flush=True)

            try:
                doc = Document(file_path)
                composer.append(doc)
                merged_files.append(file_path)

                print(f"✓ ({time.time() - t_file_start:.2f}s)")

            except Exception as e:
                print(f"✗ LỖI: {str(e)}")
                print(f"   → Khởi động lại với {len(merged_files)} file đã ghép...")

                # Save what we have so far to temp file
                temp_file = f"{output_file}.temp.docx"
                try:
                    composer.save(temp_file)

                    # Restart with temp file as master
                    master = Document(temp_file)
                    composer = Composer(master)

                    # Try again with current file
                    try:
                        doc = Document(file_path)
                        composer.append(doc)
                        merged_files.append(file_path)
                        print(f"   ✓ Thành công sau khi khởi động lại")
                    except:
                        print(f"   ✗ Vẫn thất bại, bỏ qua file này")

                    # Cleanup temp
                    try:
                        os.remove(temp_file)
                    except:
                        pass

                except Exception as e2:
                    print(f"   ✗ Không thể khởi động lại: {e2}")

        # 3. Save final result
        t_save = time.time()
        print(f"\n[Lưu file] Đang ghi file kết quả...", end=" ", flush=True)

        composer.save(output_file)

        print(f"✓ ({time.time() - t_save:.2f}s)")
        print(f"\n📊 Đã ghép thành công: {len(merged_files)}/{len(files)} file")

    except Exception as e:
        print(f"\n✗ LỖI NGHIÊM TRỌNG: {e}")
        import traceback
        traceback.print_exc()
        return False

    total_time = time.time() - total_start_time
    print(f"\n{'='*60}")
    print(f"✓ HOÀN TẤT! Tổng thời gian: {total_time:.2f}s")
    print(f"  File đầu ra: {output_file}")
    print(f"{'='*60}\n")

    return True


def _merge_composer_safe(files, output_file, total_start_time):
    """
    Method 2: Use docxcompose with safe section handling
    PROS: Fast, preserves format, handles relationships
    CONS: May fail on section conflicts
    """
    try:
        from docxcompose.composer import Composer

        # 1. Load master
        t_start = time.time()
        print(f"[1/{len(files)}] Đang tải file gốc: {os.path.basename(files[0])}...", end=" ", flush=True)

        master = Document(files[0])
        composer = Composer(master)

        print(f"✓ ({time.time() - t_start:.2f}s)")

        # 2. Merge documents
        for i, file_path in enumerate(files[1:], 2):
            t_file_start = time.time()
            print(f"[{i}/{len(files)}] Đang ghép: {os.path.basename(file_path)}...", end=" ", flush=True)

            try:
                doc = Document(file_path)
                composer.append(doc)

                print(f"✓ ({time.time() - t_file_start:.2f}s)")

            except Exception as e:
                print(f"✗ LỖI: {str(e)}")
                # Continue with remaining files
                continue

        # 3. Save
        t_save = time.time()
        print(f"\n[Lưu file] Đang ghi file kết quả...", end=" ", flush=True)

        composer.save(output_file)

        print(f"✓ ({time.time() - t_save:.2f}s)")

    except Exception as e:
        print(f"\n✗ LỖI: {e}")
        import traceback
        traceback.print_exc()
        return False

    total_time = time.time() - total_start_time
    print(f"\n{'='*60}")
    print(f"✓ HOÀN TẤT! Tổng thời gian: {total_time:.2f}s")
    print(f"  File đầu ra: {output_file}")
    print(f"{'='*60}\n")

    return True


def _merge_paragraph_api(files, output_file, total_start_time):
    """
    Method 3: Use python-docx API to copy content
    PROS: No corruption, safe, no external dependencies
    CONS: May lose some advanced formatting (tables, images)
    """
    try:
        # 1. Load master
        t_start = time.time()
        print(f"[1/{len(files)}] Đang tải file gốc: {os.path.basename(files[0])}...", end=" ", flush=True)

        master = Document(files[0])

        print(f"✓ ({time.time() - t_start:.2f}s)")

        # 2. Merge using API
        for i, file_path in enumerate(files[1:], 2):
            t_file_start = time.time()
            print(f"[{i}/{len(files)}] Đang ghép: {os.path.basename(file_path)}...", end=" ", flush=True)

            try:
                doc = Document(file_path)

                # Add page break
                master.add_page_break()

                # Copy all paragraphs
                for para in doc.paragraphs:
                    new_para = master.add_paragraph(para.text, para.style)

                    # Copy paragraph format
                    try:
                        if para.paragraph_format:
                            new_para.paragraph_format.alignment = para.paragraph_format.alignment
                            new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                            new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
                            new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
                            new_para.paragraph_format.space_before = para.paragraph_format.space_before
                            new_para.paragraph_format.space_after = para.paragraph_format.space_after
                    except:
                        pass

                    # Copy run formatting
                    new_para.clear()
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        try:
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            if run.font.size:
                                new_run.font.size = run.font.size
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.color.rgb:
                                new_run.font.color.rgb = run.font.color.rgb
                        except:
                            pass

                # Copy tables
                for table in doc.tables:
                    try:
                        new_table = master.add_table(rows=len(table.rows), cols=len(table.columns))

                        for i_row, row in enumerate(table.rows):
                            for i_col, cell in enumerate(row.cells):
                                new_table.rows[i_row].cells[i_col].text = cell.text
                    except:
                        pass

                print(f"✓ ({time.time() - t_file_start:.2f}s)")

            except Exception as e:
                print(f"✗ LỖI: {str(e)}")

        # 3. Save
        t_save = time.time()
        print(f"\n[Lưu file] Đang ghi file kết quả...", end=" ", flush=True)

        master.save(output_file)

        print(f"✓ ({time.time() - t_save:.2f}s)")

    except Exception as e:
        print(f"\n✗ LỖI: {e}")
        import traceback
        traceback.print_exc()
        return False

    total_time = time.time() - total_start_time
    print(f"\n{'='*60}")
    print(f"✓ HOÀN TẤT! Tổng thời gian: {total_time:.2f}s")
    print(f"  File đầu ra: {output_file}")
    print(f"{'='*60}\n")

    return True


# Legacy function for backward compatibility
def merge_with_doccompose(files, output_file):
    """Legacy function - uses composer_restart for best reliability"""
    return merge_docx_advanced(files, output_file, method='composer_restart')


if __name__ == "__main__":
    # Test with your files
    input_files = ["1.docx", "2.docx", "3.docx", "4.docx"]
    output_filename = "merged_output.docx"

    if not input_files:
        print("Danh sách file rỗng!")
    else:
        # Use composer_restart for best results (most reliable, no corruption)
        print("Sử dụng phương thức: composer_restart (đáng tin cậy nhất)")
        success = merge_docx_advanced(input_files, output_filename, method='composer_restart')

        if not success:
            print("\n⚠ Phương thức 1 thất bại, thử phương thức 2...")
            success = merge_docx_advanced(input_files, output_filename, method='composer_safe')

        if not success:
            print("\n⚠ Phương thức 2 thất bại, thử phương thức 3...")
            merge_docx_advanced(input_files, output_filename, method='paragraph_api')

