"""
XML Cutting - Fixed Version
Properly handles sections to prevent merge issues
"""
from docx import Document
import os
import re


def split_document_by_keyword(input_file, output_dir=None, keyword="HƯỚNG DẪN GIẢI CHI TIẾT"):
    """
    Split DOCX file into two parts based on keyword

    Args:
        input_file: Path to input DOCX file
        output_dir: Output directory (default: same as input)
        keyword: Keyword to split at

    Returns:
        tuple: (part1_path, part2_path) or (None, None) on error
    """
    print(f"\n{'='*60}")
    print(f"CHIA FILE: {os.path.basename(input_file)}")
    print(f"Từ khóa: {keyword}")
    print(f"{'='*60}\n")

    try:
        # Determine output directory
        if output_dir is None:
            output_dir = os.path.dirname(input_file)

        os.makedirs(output_dir, exist_ok=True)

        # Load document
        print("[1/5] Đang tải file...", end=" ", flush=True)
        doc = Document(input_file)
        print("✓")

        # Find split point
        print("[2/5] Tìm vị trí chia...", end=" ", flush=True)
        split_index = -1

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()
            if keyword.upper() in text:
                split_index = i
                break

        if split_index == -1:
            print(f"✗\n❌ Không tìm thấy từ khóa '{keyword}'")
            return None, None

        print(f"✓ (tại đoạn {split_index})")

        # Create part 1 (before keyword)
        print("[3/5] Tạo phần 1 (trước từ khóa)...", end=" ", flush=True)

        doc1 = Document()

        # Copy paragraphs before split point
        for i in range(split_index):
            para = doc.paragraphs[i]
            new_para = doc1.add_paragraph(para.text, para.style)

            # Copy formatting
            new_para.paragraph_format.alignment = para.paragraph_format.alignment

        # Save part 1
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        part1_path = os.path.join(output_dir, f"DE_{base_name}.docx")
        doc1.save(part1_path)

        print(f"✓ ({split_index} đoạn)")

        # Create part 2 (from keyword onwards)
        print("[4/5] Tạo phần 2 (từ từ khóa)...", end=" ", flush=True)

        doc2 = Document()

        # Copy paragraphs from split point onwards
        for i in range(split_index, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            new_para = doc2.add_paragraph(para.text, para.style)

            # Copy formatting
            new_para.paragraph_format.alignment = para.paragraph_format.alignment

        # Save part 2
        part2_path = os.path.join(output_dir, f"GIAI_{base_name}.docx")
        doc2.save(part2_path)

        print(f"✓ ({len(doc.paragraphs) - split_index} đoạn)")

        # Summary
        print("[5/5] Hoàn tất!")
        print(f"\n{'='*60}")
        print(f"✓ THÀNH CÔNG!")
        print(f"  • Phần 1: {part1_path}")
        print(f"  • Phần 2: {part2_path}")
        print(f"{'='*60}\n")

        return part1_path, part2_path

    except Exception as e:
        print(f"✗\n❌ LỖI: {e}")
        import traceback
        traceback.print_exc()
        return None, None


def clean_document_headers_footers(doc_path, output_path=None):
    """
    Remove headers, footers, and page numbers from document

    Args:
        doc_path: Path to input document
        output_path: Path to output (default: overwrite input)
    """
    if output_path is None:
        output_path = doc_path

    print(f"Dọn dẹp headers/footers: {os.path.basename(doc_path)}...", end=" ", flush=True)

    try:
        doc = Document(doc_path)

        # Remove headers and footers from all sections
        for section in doc.sections:
            # Clear headers
            section.header.is_linked_to_previous = False
            for para in section.header.paragraphs:
                para.clear()

            # Clear footers
            section.footer.is_linked_to_previous = False
            for para in section.footer.paragraphs:
                para.clear()

        doc.save(output_path)
        print("✓")

    except Exception as e:
        print(f"✗ LỖI: {e}")


if __name__ == "__main__":
    import sys

    # Check command line arguments
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        keyword = sys.argv[2] if len(sys.argv) > 2 else "HƯỚNG DẪN GIẢI CHI TIẾT"
    else:
        # Default test
        input_file = "test.docx"
        keyword = "HƯỚNG DẪN GIẢI CHI TIẾT"

        print("Cách dùng: python xml_cutting_clean.py <file.docx> [từ_khóa]")
        print(f"Sử dụng giá trị mặc định...\n")

    # Check if file exists
    if not os.path.exists(input_file):
        print(f"❌ Không tìm thấy file: {input_file}")
        sys.exit(1)

    # Split document
    part1, part2 = split_document_by_keyword(input_file, keyword=keyword)

    if part1 and part2:
        # Clean headers/footers
        print("\nDọn dẹp headers và footers...")
        clean_document_headers_footers(part1)
        clean_document_headers_footers(part2)

        print("\n✅ Hoàn tất!")
        print(f"   File đầu ra:")
        print(f"   - {part1}")
        print(f"   - {part2}")
    else:
        print("\n❌ Thất bại!")
        sys.exit(1)

