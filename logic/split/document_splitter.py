"""
Document Split Worker
Handles document splitting operations in background thread
"""
from docx import Document
from docx.shared import Pt, Inches
import os
import copy
import re
from datetime import datetime
from typing import List, Optional, Dict, Any


def split_documents_worker_simple(thread, config: dict):
    """
    Simple worker function for splitting documents (single thread processes assigned files)

    This is designed to work with ThreadManager.create_worker_threads() which
    distributes files among multiple threads.

    Args:
        thread: WorkerThread instance (automatically passed)
        config: Split configuration dict with:
            - mode: 'keyword', 'half', or 'custom'
            - files: List of file paths assigned to THIS thread
            - output_dir: Output directory path
            - options: Dict with remove_header, remove_footer, remove_page_numbers, remove_first_shape
            - keyword: Keyword for keyword mode
            - part_selection: Dict with save_de, save_giai
            - ranges: List of SplitRange objects (for custom mode)

    Returns:
        Dict with results
    """
    try:
        mode = config.get('mode', 'keyword')
        files = config.get('files', [])
        output_dir = config.get('output_dir', '')
        options = config.get('options', {})
        ranges = config.get('ranges', [])
        keyword = config.get('keyword', 'HƯỚNG DẪN GIẢI CHI TIẾT')
        part_selection = config.get('part_selection', {'save_de': True, 'save_giai': True})

        thread.log(f"Worker started with {len(files)} files assigned", "INFO")
        thread.log(f"Mode: {mode}", "INFO")
        thread.current_task = f"Splitting ({mode})"

        total_files = len(files)
        output_files = []

        for i, file_path in enumerate(files):
            # Check for pause
            thread.wait_if_paused()

            if thread.should_stop():
                thread.log("Split operation stopped by user", "WARNING")
                return {
                    'output_files': output_files,
                    'stopped': True,
                    'processed': i
                }

            # Update current file and mark as processing
            file_name = os.path.basename(file_path)
            thread.current_file = file_path
            thread.log(f"[{i+1}/{total_files}] Processing: {file_name}", "INFO")

            # Track file start time
            file_start_time = datetime.now()

            # Update main progress
            progress = int((i / total_files) * 100)
            thread.update_progress(progress, f"Splitting ({mode})")

            try:
                # Process based on mode
                if mode == 'keyword':
                    result_files = _split_by_keyword(thread, file_path, output_dir, options, keyword, part_selection)
                elif mode == 'half':
                    result_files = _split_in_half(thread, file_path, output_dir, options)
                else:  # custom
                    result_files = _split_by_ranges(thread, file_path, output_dir, options, ranges)

                output_files.extend(result_files)
                thread.processed_files += 1

                # Calculate file processing time
                file_elapsed = (datetime.now() - file_start_time).total_seconds()
                thread.mark_file_completed(file_path, file_elapsed)

                thread.log(f"✓ Successfully processed {file_name} in {file_elapsed:.2f}s", "SUCCESS")

            except Exception as e:
                error_msg = str(e)
                thread.log(f"✗ Error processing {file_name}: {error_msg}", "ERROR")
                import traceback
                thread.log(f"  {traceback.format_exc()}", "ERROR")

                # Mark file as error
                thread.mark_file_error(file_path, error_msg)

        # Final progress
        thread.current_file = None
        thread.update_progress(100, "Completed")
        thread.log(f"Worker completed: {thread.processed_files}/{total_files} files processed", "SUCCESS")

        return {
            'output_files': output_files,
            'stopped': False,
            'processed': total_files
        }

    except Exception as e:
        thread.log(f"Fatal error: {str(e)}", "ERROR")
        import traceback
        thread.log(traceback.format_exc(), "ERROR")
        return {
            'output_files': [],
            'stopped': True,
            'error': str(e)
        }


def split_documents_worker(thread, config: dict):
    """
    Worker function for splitting documents

    Uses multi-threading for batch processing when more than 2 files.

    Args:
        thread: WorkerThread instance (automatically passed)
        config: Split configuration dict with:
            - mode: 'keyword', 'half', or 'custom'
            - files: List of file paths to split
            - output_dir: Output directory path
            - options: Dict with remove_header, remove_footer, remove_page_numbers, remove_first_shape
            - keyword: Keyword for keyword mode
            - part_selection: Dict with save_de, save_giai
            - ranges: List of SplitRange objects (for custom mode)

    Returns:
        Dict with results
    """
    try:
        mode = config.get('mode', 'keyword')
        files = config.get('files', [])
        output_dir = config.get('output_dir', '')
        options = config.get('options', {})
        ranges = config.get('ranges', [])
        keyword = config.get('keyword', 'HƯỚNG DẪN GIẢI CHI TIẾT')
        part_selection = config.get('part_selection', {'save_de': True, 'save_giai': True})

        thread.log("Starting document split operation...", "INFO")
        thread.log(f"Mode: {mode}", "INFO")
        thread.log(f"Files to process: {len(files)}", "INFO")
        thread.log(f"Output directory: {output_dir}", "INFO")

        total_files = len(files)

        # Multi-threading for batch processing
        if total_files > 2:
            thread.log(f"Using multi-threading for {total_files} files...", "INFO")
            return _process_files_multithreaded(thread, files, mode, output_dir, options,
                                               keyword, part_selection, ranges)
        else:
            thread.log("Processing files sequentially...", "INFO")
            return _process_files_sequential(thread, files, mode, output_dir, options,
                                            keyword, part_selection, ranges)

    except Exception as e:
        thread.log(f"Fatal error in split worker: {str(e)}", "ERROR")
        import traceback
        thread.log(traceback.format_exc(), "ERROR")
        return {
            'output_files': [],
            'stopped': True,
            'error': str(e)
        }


def _process_files_sequential(thread, files, mode, output_dir, options, keyword, part_selection, ranges):
    """Process files sequentially (for 1-2 files)"""
    thread.update_progress(0, "Starting split operation...")

    total_files = len(files)
    output_files = []

    for i, file_path in enumerate(files):
        if thread.should_stop():
            thread.log("Split operation stopped by user", "WARNING")
            return {
                'output_files': output_files,
                'stopped': True,
                'processed': i
            }

        file_name = os.path.basename(file_path)
        thread.log(f"\n[{i+1}/{total_files}] Processing: {file_name}", "INFO")
        thread.update_progress(
            int((i / total_files) * 100),
            f"Processing {file_name}..."
        )

        try:
            if mode == 'keyword':
                result_files = _split_by_keyword(thread, file_path, output_dir, options, keyword, part_selection)
            elif mode == 'half':
                result_files = _split_in_half(thread, file_path, output_dir, options)
            else:  # custom
                result_files = _split_by_ranges(thread, file_path, output_dir, options, ranges)

            output_files.extend(result_files)
            thread.log(f"  ✓ Successfully processed {file_name}", "SUCCESS")

        except Exception as e:
            thread.log(f"  ✗ Error processing {file_name}: {str(e)}", "ERROR")
            import traceback
            thread.log(f"  {traceback.format_exc()}", "ERROR")

    # Final progress
    thread.update_progress(100, "Split operation completed")

    if not thread.should_stop():
        thread.log("\n=== Split operation completed ===", "SUCCESS")
        thread.log(f"Processed {total_files} file(s)", "SUCCESS")
        thread.log(f"Created {len(output_files)} output file(s)", "SUCCESS")
        thread.log(f"Output directory: {output_dir}", "INFO")

    return {
        'output_files': output_files,
        'stopped': False,
        'processed': total_files
    }


def _process_files_multithreaded(thread, files, mode, output_dir, options, keyword, part_selection, ranges):
    """Process files using multiple threads (for 3+ files)"""
    import threading
    import queue
    from concurrent.futures import ThreadPoolExecutor, as_completed

    thread.update_progress(0, "Starting multi-threaded split operation...")

    total_files = len(files)
    output_files = []
    processed_count = 0
    lock = threading.Lock()

    # Determine number of worker threads (max 4 or number of files, whichever is smaller)
    num_workers = min(4, total_files)
    thread.log(f"Using {num_workers} worker threads", "INFO")

    def process_single_file(file_path, file_index):
        """Process a single file (runs in thread pool)"""
        file_name = os.path.basename(file_path)

        # Create a simple logger that's thread-safe
        class SimpleLogger:
            def log(self, msg, level="INFO"):
                with lock:
                    thread.log(f"[Worker {file_index}] {msg}", level)

        logger = SimpleLogger()

        try:
            logger.log(f"Processing: {file_name}", "INFO")

            if mode == 'keyword':
                result_files = _split_by_keyword(logger, file_path, output_dir, options, keyword, part_selection)
            elif mode == 'half':
                result_files = _split_in_half(logger, file_path, output_dir, options)
            else:  # custom
                result_files = _split_by_ranges(logger, file_path, output_dir, options, ranges)

            logger.log(f"✓ Successfully processed {file_name}", "SUCCESS")
            return {'success': True, 'files': result_files, 'name': file_name}

        except Exception as e:
            logger.log(f"✗ Error processing {file_name}: {str(e)}", "ERROR")
            return {'success': False, 'files': [], 'name': file_name, 'error': str(e)}

    # Process files in parallel
    with ThreadPoolExecutor(max_workers=num_workers) as executor:
        # Submit all tasks
        future_to_file = {
            executor.submit(process_single_file, file_path, i): (file_path, i)
            for i, file_path in enumerate(files, 1)
        }

        # Process completed tasks
        for future in as_completed(future_to_file):
            if thread.should_stop():
                thread.log("Split operation stopped by user", "WARNING")
                executor.shutdown(wait=False)
                return {
                    'output_files': output_files,
                    'stopped': True,
                    'processed': processed_count
                }

            result = future.result()
            processed_count += 1

            if result['success']:
                output_files.extend(result['files'])

            # Update progress
            progress = int((processed_count / total_files) * 100)
            thread.update_progress(progress, f"Processed {processed_count}/{total_files} files")

    # Final progress
    thread.update_progress(100, "Multi-threaded split completed")

    if not thread.should_stop():
        thread.log("\n=== Multi-threaded split operation completed ===", "SUCCESS")
        thread.log(f"Processed {total_files} file(s) using {num_workers} threads", "SUCCESS")
        thread.log(f"Created {len(output_files)} output file(s)", "SUCCESS")
        thread.log(f"Output directory: {output_dir}", "INFO")

    return {
        'output_files': output_files,
        'stopped': False,
        'processed': total_files,
        'threads_used': num_workers
    }


def _clean_first_paragraph_safely(thread, doc: Document):
    """
    Safely clean first paragraph - new improved version to prevent merge errors.

    Strategy:
    1. Check if paragraph contains <w:sectPr> - if so, keep paragraph structure
    2. Remove embedded objects (images, shapes) by clearing relationship IDs
    3. Clear text content from runs but keep run structure
    4. This prevents "list index out of range" errors during merge

    This is the new method that replaces _clean_first_paragraph_deeply for better
    compatibility with docxcompose merging.
    """
    body = doc.element.body

    if len(body) == 0:
        return

    first_p = body[0]

    # Only process paragraph elements
    if not first_p.tag.endswith('p'):
        return

    xml_str = first_p.xml

    # Remove relationship IDs for images/shapes first
    rids = re.findall(r'(?:r:id|r:embed)="([^"]+)"', xml_str)
    if rids:
        thread.log(f"    [-] Removing {len(rids)} embedded objects from first paragraph...", "INFO")
        for rid in rids:
            try:
                if rid in doc.part.rels:
                    doc.part.rels.pop(rid)
            except Exception as e:
                pass  # Continue if relationship doesn't exist

    # If paragraph contains sectPr, keep it but clean content carefully
    has_sectp = "w:sectPr" in xml_str

    if has_sectp:
        thread.log("    [!] First paragraph contains <w:sectPr>, preserving section settings.", "INFO")

        # Keep sectPr, clean everything else
        children = list(first_p)
        for child in children:
            if child.tag.endswith("sectPr"):
                continue  # Preserve section properties

            if child.tag.endswith("pPr"):
                # Keep paragraph properties but remove run properties if any
                ppr_children = list(child)
                for ppr_child in ppr_children:
                    if not ppr_child.tag.endswith("sectPr"):
                        child.remove(ppr_child)
                continue

            # For w:r (run) elements: clear content but keep structure
            if child.tag.endswith("r"):
                # Keep the run element, but clear its content children
                run_children = list(child)
                for run_child in run_children:
                    # Remove text, drawing, picture elements
                    if (run_child.tag.endswith("t") or
                        run_child.tag.endswith("drawing") or
                        run_child.tag.endswith("pict")):
                        child.remove(run_child)
            else:
                # Remove other elements
                first_p.remove(child)

        thread.log("    [✓] Safely cleaned first paragraph (preserved structure for merging).", "SUCCESS")
        return

    # Normal case: paragraph doesn't contain sectPr, safe to remove entirely
    body.remove(first_p)
    thread.log("    [✓] Removed first paragraph (no sectPr).", "SUCCESS")


def _split_by_keyword(thread, file_path: str, output_dir: str, options: dict,
                      keyword: str, part_selection: dict) -> List[str]:
    """
    Split document by keyword (DE/GIAI mode) - follows xml_cutting.py logic

    Creates:
    - DE_<name>.docx: Content before keyword (not including keyword)
    - GIAI_<name>.docx: Content from keyword to end
    """
    import time
    operation_start = time.time()

    thread.log(f"  Splitting by keyword: '{keyword}'...", "INFO")
    thread.update_sub_progress(0, "Loading document")

    # Load document
    load_start = time.time()
    doc = Document(file_path)
    load_time = time.time() - load_start
    thread.log(f"    [1/5] Document loaded in {load_time:.3f}s", "INFO")
    thread.update_sub_progress(20, "Document loaded")

    # Log source document margins
    if len(doc.sections) > 0:
        def _emu_to_cm(value):
            """Convert EMU to cm for display"""
            try:
                inches = value / 914400
                cm = inches * 2.54
                return f"{cm:.2f}cm"
            except:
                return str(value)

        src_section = doc.sections[0]
        thread.log("    [2/5] Source document page settings:", "INFO")
        thread.log(f"           Top Margin: {_emu_to_cm(src_section.top_margin)}", "INFO")
        thread.log(f"           Bottom Margin: {_emu_to_cm(src_section.bottom_margin)}", "INFO")
        thread.log(f"           Left Margin: {_emu_to_cm(src_section.left_margin)}", "INFO")
        thread.log(f"           Right Margin: {_emu_to_cm(src_section.right_margin)}", "INFO")
        thread.log(f"           Page Size: {_emu_to_cm(src_section.page_width)} x {_emu_to_cm(src_section.page_height)}", "INFO")
        thread.log(f"           Total Sections: {len(doc.sections)}", "INFO")

    thread.update_sub_progress(30, "Analyzing structure")

    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output_files = []

    thread.update_sub_progress(40, "Creating split files")

    # Create DE file (before keyword)
    if part_selection.get('save_de', True):
        de_start = time.time()
        thread.log("    [3/5] Creating DE file (before keyword)...", "INFO")
        doc_de = copy.deepcopy(doc)

        # Clean first paragraph using new safer method
        if options.get('remove_first_shape', True):
            _clean_first_paragraph_safely(thread, doc_de)

        # Remove content from keyword onwards
        body_de = doc_de.element.body
        found_keyword = False
        to_delete = []

        for element in body_de.iterchildren():
            if found_keyword:
                to_delete.append(element)
                continue

            if element.tag.endswith('p') or element.tag.endswith('tbl'):
                text = _get_full_text(element).upper()
                if keyword.upper() in text:
                    found_keyword = True
                    # Delete the keyword line itself
                    to_delete.append(element)

        if not found_keyword:
            thread.log(f"      ⚠ Keyword '{keyword}' not found. DE file will contain full document.", "WARNING")

        for el in to_delete:
            body_de.remove(el)

        # CRITICAL: Ensure document has at least one section with proper sectPr
        # Word documents need sectPr either at body level OR in the last paragraph
        # Strategy: Add sectPr as LAST child of body (after all content)
        from lxml import etree
        from docx.oxml import OxmlElement
        import copy as copy_module

        # Get original section properties
        original_sect = doc.element.body.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')

        if original_sect is not None:
            # Remove any existing sectPr from body or paragraphs
            for existing_sect in body_de.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr'):
                existing_sect.getparent().remove(existing_sect)

            # Ensure at least one paragraph exists
            paragraphs = body_de.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
            if len(paragraphs) == 0:
                new_p = OxmlElement('w:p')
                body_de.append(new_p)
                thread.log("      ⚠ No paragraphs found - created one", "WARNING")

            # Clone sectPr from source and add as LAST child of body
            new_sect = copy_module.deepcopy(original_sect)
            body_de.append(new_sect)
            thread.log("      ✓ Restored section properties to DE file (body-level)", "INFO")
        else:
            thread.log("      ⚠ No section properties found in original document", "WARNING")

        # Apply options (remove headers, footers, etc.)
        _apply_options(thread, doc_de, options)

        # Verify we have sections before copying settings
        if len(doc_de.sections) == 0:
            thread.log("      ❌ ERROR: DE file still has no sections after restoration!", "ERROR")
        else:
            thread.log(f"      ✓ DE file has {len(doc_de.sections)} section(s)", "INFO")

        # ALWAYS copy page settings from source AFTER applying options to ensure margins are preserved
        _copy_page_settings(doc, doc_de, thread)
        thread.log("      ✓ Page margins copied from source to DE file", "SUCCESS")

        # Save
        output_de = os.path.join(output_dir, f"DE_{base_name}.docx")
        doc_de.save(output_de)
        output_files.append(output_de)
        de_time = time.time() - de_start
        thread.log(f"      ✓ Saved: {os.path.basename(output_de)} ({de_time:.3f}s)", "SUCCESS")
        thread.update_sub_progress(60, "DE file created")

    # Create GIAI file (from keyword to end)
    if part_selection.get('save_giai', True):
        giai_start = time.time()
        thread.log("    [4/5] Creating GIAI file (from keyword to end)...", "INFO")
        doc_giai = copy.deepcopy(doc)

        # Clean first paragraph using new safer method
        if options.get('remove_first_shape', True):
            _clean_first_paragraph_safely(thread, doc_giai)

        # Find the keyword and keep header + content from keyword
        body_giai = doc_giai.element.body
        children = list(body_giai.iterchildren())

        start_delete_idx = -1
        end_delete_idx = -1
        header_idx = -1

        for i, element in enumerate(children):
            text = _get_full_text(element).upper()

            # Find header (BÀI: or ĐỀ TEST)
            if header_idx == -1 and ("BÀI:" in text or "ĐỀ TEST" in text):
                header_idx = i
                start_delete_idx = i + 1  # Delete from after header

            # Find keyword
            if keyword.upper() in text:
                end_delete_idx = i  # Delete up to (but not including) keyword
                break

        if start_delete_idx > 0 and end_delete_idx > start_delete_idx:
            thread.log(f"      Removing content from index {start_delete_idx} to {end_delete_idx - 1}", "INFO")
            nodes_to_remove = children[start_delete_idx:end_delete_idx]
            for node in nodes_to_remove:
                body_giai.remove(node)
        elif end_delete_idx == -1:
            thread.log(f"      ⚠ Keyword '{keyword}' not found. GIAI file may contain unwanted content.", "WARNING")

        # Apply options (remove headers, footers, etc.)
        _apply_options(thread, doc_giai, options)

        # ALWAYS copy page settings from source AFTER applying options to ensure margins are preserved
        _copy_page_settings(doc, doc_giai, thread)
        thread.log("      ✓ Page margins copied from source to GIAI file", "SUCCESS")

        # Save
        output_giai = os.path.join(output_dir, f"GIAI_{base_name}.docx")
        doc_giai.save(output_giai)
        output_files.append(output_giai)
        giai_time = time.time() - giai_start
        thread.log(f"      ✓ Saved: {os.path.basename(output_giai)} ({giai_time:.3f}s)", "SUCCESS")
        thread.update_sub_progress(80, "GIAI file created")

    # Final timing
    operation_time = time.time() - operation_start
    thread.log(f"    [5/5] Split operation completed in {operation_time:.3f}s", "SUCCESS")
    thread.update_sub_progress(100, "Split complete")

    return output_files


def _get_full_text(element) -> str:
    """Get all text from an element (helper for keyword detection)"""
    texts = []
    for node in element.iter():
        if node.text:
            texts.append(node.text)
    return "".join(texts)


def _split_in_half(thread, file_path: str, output_dir: str, options: dict) -> List[str]:
    """Split document in half"""
    thread.log("  Splitting document in half...", "INFO")

    # Load document
    doc = Document(file_path)

    # Log source document margins
    if len(doc.sections) > 0:
        def _emu_to_cm(value):
            """Convert EMU to cm for display"""
            try:
                inches = value / 914400
                cm = inches * 2.54
                return f"{cm:.2f}cm"
            except:
                return str(value)

        src_section = doc.sections[0]
        thread.log("  📖 SOURCE DOCUMENT PAGE SETTINGS:", "INFO")
        thread.log(f"     Top Margin: {_emu_to_cm(src_section.top_margin)}", "INFO")
        thread.log(f"     Bottom Margin: {_emu_to_cm(src_section.bottom_margin)}", "INFO")
        thread.log(f"     Left Margin: {_emu_to_cm(src_section.left_margin)}", "INFO")
        thread.log(f"     Right Margin: {_emu_to_cm(src_section.right_margin)}", "INFO")
        thread.log(f"     Page Size: {_emu_to_cm(src_section.page_width)} x {_emu_to_cm(src_section.page_height)}", "INFO")

    # Clean first paragraph using new safer method
    if options.get('remove_first_shape', True):
        _clean_first_paragraph_safely(thread, doc)

    # Get all paragraphs and tables
    body_elements = list(doc.element.body)
    total_elements = len(body_elements)
    mid_point = total_elements // 2

    thread.log(f"    Total elements: {total_elements}", "INFO")
    thread.log(f"    Split point: {mid_point}", "INFO")

    output_files = []

    # Create first half
    doc1 = copy.deepcopy(doc)
    _remove_elements_after_index(doc1, mid_point)
    _apply_options(thread, doc1, options)

    # ALWAYS copy page settings AFTER applying options to ensure margins are preserved
    _copy_page_settings(doc, doc1, thread)
    thread.log("    ✓ Page margins copied to first half", "SUCCESS")

    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output1 = os.path.join(output_dir, f"{base_name}_part1.docx")
    doc1.save(output1)
    output_files.append(output1)
    thread.log(f"    ✓ Saved: {os.path.basename(output1)}", "SUCCESS")

    # Create second half
    doc2 = copy.deepcopy(doc)
    _remove_elements_before_index(doc2, mid_point)
    _apply_options(thread, doc2, options)

    # ALWAYS copy page settings AFTER applying options to ensure margins are preserved
    _copy_page_settings(doc, doc2, thread)
    thread.log("    ✓ Page margins copied to second half", "SUCCESS")

    output2 = os.path.join(output_dir, f"{base_name}_part2.docx")
    doc2.save(output2)
    output_files.append(output2)
    thread.log(f"    ✓ Saved: {os.path.basename(output2)}", "SUCCESS")

    return output_files


def _split_by_ranges(thread, file_path: str, output_dir: str, options: dict, ranges: List) -> List[str]:
    """Split document by custom ranges"""
    thread.log(f"  Splitting by {len(ranges)} custom range(s)...", "INFO")

    # Load document
    doc = Document(file_path)

    # Log source document margins
    if len(doc.sections) > 0:
        def _emu_to_cm(value):
            """Convert EMU to cm for display"""
            try:
                inches = value / 914400
                cm = inches * 2.54
                return f"{cm:.2f}cm"
            except:
                return str(value)

        src_section = doc.sections[0]
        thread.log("  📖 SOURCE DOCUMENT PAGE SETTINGS:", "INFO")
        thread.log(f"     Top Margin: {_emu_to_cm(src_section.top_margin)}", "INFO")
        thread.log(f"     Bottom Margin: {_emu_to_cm(src_section.bottom_margin)}", "INFO")
        thread.log(f"     Left Margin: {_emu_to_cm(src_section.left_margin)}", "INFO")
        thread.log(f"     Right Margin: {_emu_to_cm(src_section.right_margin)}", "INFO")
        thread.log(f"     Page Size: {_emu_to_cm(src_section.page_width)} x {_emu_to_cm(src_section.page_height)}", "INFO")

    # Clean first paragraph using new safer method
    if options.get('remove_first_shape', True):
        _clean_first_paragraph_safely(thread, doc)

    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output_files = []

    for i, split_range in enumerate(ranges, 1):
        thread.log(f"    Processing range #{i}...", "INFO")

        try:
            # Create a copy for this range
            range_doc = copy.deepcopy(doc)

            # Extract the range
            _extract_range(thread, range_doc, split_range)

            # Apply options (remove headers, footers, etc.)
            _apply_options(thread, range_doc, options)

            # ALWAYS copy page settings from source AFTER applying options to ensure margins are preserved
            _copy_page_settings(doc, range_doc, thread)
            thread.log(f"      ✓ Page margins copied to range #{i}", "SUCCESS")

            # Save
            output_path = os.path.join(output_dir, f"{base_name}_range{i}.docx")
            range_doc.save(output_path)
            output_files.append(output_path)
            thread.log(f"      ✓ Saved: {os.path.basename(output_path)}", "SUCCESS")

        except Exception as e:
            thread.log(f"      ✗ Error processing range #{i}: {str(e)}", "ERROR")

    return output_files


def _extract_range(thread, doc: Document, split_range):
    """Extract a specific range from document"""
    body_elements = list(doc.element.body)

    # Find start index
    start_idx = 0
    if not split_range.from_start:
        start_idx = _find_keyword_index(body_elements, split_range.from_keyword,
                                       split_range.from_include_all)
        if start_idx == -1:
            raise ValueError(f"Start keyword '{split_range.from_keyword}' not found")

    # Find end index
    end_idx = len(body_elements)
    if not split_range.to_end:
        end_idx = _find_keyword_index(body_elements, split_range.to_keyword,
                                     split_range.to_include_all, start_from=start_idx)
        if end_idx == -1:
            raise ValueError(f"End keyword '{split_range.to_keyword}' not found")

    thread.log(f"      Range indices: {start_idx} to {end_idx}", "INFO")

    # Remove elements outside the range
    _remove_elements_before_index(doc, start_idx)
    _remove_elements_after_index(doc, end_idx - start_idx)


def _find_keyword_index(elements: List, keyword: str, include_all: bool = False,
                       start_from: int = 0) -> int:
    """Find index of element containing keyword"""
    keyword_upper = keyword.upper()

    for i in range(start_from, len(elements)):
        element = elements[i]
        text = _get_element_text(element).upper()

        if keyword_upper in text:
            if include_all:
                # Find first occurrence with this keyword
                return i
            else:
                # Return this specific element
                return i

    return -1


def _get_element_text(element) -> str:
    """Get all text from an element"""
    texts = []
    for node in element.iter():
        if node.text:
            texts.append(node.text)
    return "".join(texts)


def _copy_page_settings(source_doc: Document, target_doc: Document, thread=None):
    """
    Copy page settings from source document to target document.
    This ensures all split files have the same page settings for proper merging.

    Copies:
    - Page size (width, height)
    - Margins (top, bottom, left, right)
    - Orientation
    - Header/footer distance
    """
    try:
        # Get first section from source
        if len(source_doc.sections) == 0 or len(target_doc.sections) == 0:
            if thread:
                thread.log("      ⚠ No sections found in source or target document", "WARNING")
            return

        source_section = source_doc.sections[0]

        # Log source margin values
        if thread:
            def _mm_to_cm(value):
                """Convert EMU to cm for display"""
                try:
                    # python-docx returns values in EMUs (English Metric Units)
                    # 1 inch = 914400 EMUs, 1 inch = 2.54 cm
                    inches = value / 914400
                    cm = inches * 2.54
                    return f"{cm:.2f}cm"
                except:
                    return str(value)

            thread.log("      📏 SOURCE Document Margins:", "INFO")
            thread.log(f"         Top: {_mm_to_cm(source_section.top_margin)}", "INFO")
            thread.log(f"         Bottom: {_mm_to_cm(source_section.bottom_margin)}", "INFO")
            thread.log(f"         Left: {_mm_to_cm(source_section.left_margin)}", "INFO")
            thread.log(f"         Right: {_mm_to_cm(source_section.right_margin)}", "INFO")
            thread.log(f"         Page Size: {_mm_to_cm(source_section.page_width)} x {_mm_to_cm(source_section.page_height)}", "INFO")

        # Apply to all sections in target
        for idx, target_section in enumerate(target_doc.sections):
            # Log BEFORE copy
            if thread:
                thread.log(f"      📄 TARGET Section #{idx+1} BEFORE copy:", "DEBUG")
                thread.log(f"         Top: {_mm_to_cm(target_section.top_margin)}", "DEBUG")
                thread.log(f"         Bottom: {_mm_to_cm(target_section.bottom_margin)}", "DEBUG")
                thread.log(f"         Left: {_mm_to_cm(target_section.left_margin)}", "DEBUG")
                thread.log(f"         Right: {_mm_to_cm(target_section.right_margin)}", "DEBUG")

            # Copy page size
            target_section.page_width = source_section.page_width
            target_section.page_height = source_section.page_height

            # Copy orientation
            target_section.orientation = source_section.orientation

            # Copy margins
            target_section.top_margin = source_section.top_margin
            target_section.bottom_margin = source_section.bottom_margin
            target_section.left_margin = source_section.left_margin
            target_section.right_margin = source_section.right_margin

            # Copy header/footer settings
            target_section.header_distance = source_section.header_distance
            target_section.footer_distance = source_section.footer_distance

            # Log AFTER copy
            if thread:
                thread.log(f"      ✅ TARGET Section #{idx+1} AFTER copy:", "INFO")
                thread.log(f"         Top: {_mm_to_cm(target_section.top_margin)}", "INFO")
                thread.log(f"         Bottom: {_mm_to_cm(target_section.bottom_margin)}", "INFO")
                thread.log(f"         Left: {_mm_to_cm(target_section.left_margin)}", "INFO")
                thread.log(f"         Right: {_mm_to_cm(target_section.right_margin)}", "INFO")

    except Exception as e:
        # Non-critical error, but log it for debugging
        import traceback
        error_msg = f"Warning: Could not copy page settings: {str(e)}"
        if thread:
            thread.log(f"      ❌ {error_msg}", "ERROR")
            thread.log(f"      {traceback.format_exc()}", "ERROR")
        else:
            print(error_msg)
            print(traceback.format_exc())


def _remove_elements_after_index(doc: Document, index: int):
    """Remove all elements after given index"""
    body = doc.element.body
    elements = list(body)

    # Keep sectPr if exists in last element
    for i in range(index, len(elements)):
        element = elements[i]

        # Check if this element has sectPr - if so, preserve it
        if element.tag.endswith('p') and "w:sectPr" in element.xml:
            # Keep sectPr, remove other content
            children = list(element)
            for child in children:
                if not child.tag.endswith("sectPr"):
                    element.remove(child)
        else:
            body.remove(element)


def _remove_elements_before_index(doc: Document, index: int):
    """Remove all elements before given index"""
    body = doc.element.body
    elements = list(body)

    for i in range(index):
        if i < len(elements):
            body.remove(elements[i])


def _apply_options(thread, doc: Document, options: dict):
    """Apply split options to document"""
    if options.get('remove_header'):
        _remove_headers(thread, doc)

    if options.get('remove_footer'):
        _remove_footers(thread, doc)

    if options.get('remove_page_numbers'):
        _remove_page_numbers(thread, doc)


def _remove_headers(thread, doc: Document):
    """Remove all headers from document - complete removal"""
    try:
        for section in doc.sections:
            # Get header element
            header = section.header
            header.is_linked_to_previous = False

            # Get the header XML element
            header_element = header._element

            # Remove all children from header
            for child in list(header_element):
                header_element.remove(child)

            # Also try alternative method - clear all paragraphs
            for paragraph in list(header.paragraphs):
                # Remove all runs from paragraph
                for run in list(paragraph.runs):
                    run._element.getparent().remove(run._element)
                # Clear paragraph
                p_element = paragraph._element
                p_element.getparent().remove(p_element)

            thread.log(f"    Removed header from section", "INFO")
    except Exception as e:
        thread.log(f"    Warning: Could not completely remove headers: {str(e)}", "WARNING")


def _remove_footers(thread, doc: Document):
    """Remove all footers from document - complete removal"""
    try:
        for section in doc.sections:
            # Get footer element
            footer = section.footer
            footer.is_linked_to_previous = False

            # Get the footer XML element
            footer_element = footer._element

            # Remove all children from footer
            for child in list(footer_element):
                footer_element.remove(child)

            # Also try alternative method - clear all paragraphs
            for paragraph in list(footer.paragraphs):
                # Remove all runs from paragraph
                for run in list(paragraph.runs):
                    run._element.getparent().remove(run._element)
                # Clear paragraph
                p_element = paragraph._element
                p_element.getparent().remove(p_element)

            thread.log(f"    Removed footer from section", "INFO")
    except Exception as e:
        thread.log(f"    Warning: Could not completely remove footers: {str(e)}", "WARNING")


def _remove_page_numbers(thread, doc: Document):
    """Remove page numbers from document"""
    # Page numbers are typically in footers
    # This is a simplified version - full implementation would need to
    # parse the XML and remove specific page number fields
    try:
        for section in doc.sections:
            # Remove from footers
            for paragraph in section.footer.paragraphs:
                # Check if paragraph contains page number field
                if 'PAGE' in paragraph.text or paragraph.text.strip().isdigit():
                    paragraph.clear()
    except Exception as e:
        thread.log(f"    Warning: Could not remove page numbers: {str(e)}", "WARNING")
