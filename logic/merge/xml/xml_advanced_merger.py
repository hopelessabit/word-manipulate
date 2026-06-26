# -*- coding: utf-8 -*-
"""
xml_advanced_merger.py
----------------------
Advanced XML-based document merger with conflict resolution.

Features:
- Pre-merge validation of DOCX files
- Unique ID generation for styles, numbering, and relationships
- Conflict detection and resolution
- Robust error handling
- Progress tracking and detailed logging
"""

import os
import time
import zipfile
import xml.etree.ElementTree as ET
from typing import List, Optional, Callable, Dict, Tuple
from docx import Document


# Word XML namespaces
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'rel': 'http://schemas.openxmlformats.org/package/2006/relationships'
}

for prefix, uri in NAMESPACES.items():
    ET.register_namespace(prefix, uri)


class ValidationError(Exception):
    """Custom exception for validation errors"""
    pass


class XMLAdvancedMerger:
    """
    Advanced XML-based merger with conflict resolution.
    Ensures unique IDs for all formatting elements to prevent conflicts.
    """

    def __init__(self, progress_callback: Optional[Callable] = None,
                 log_callback: Optional[Callable] = None):
        """
        Initialize the advanced XML merger.

        Args:
            progress_callback: Callback function(progress: float, message: str)
            log_callback: Callback function(message: str, level: str)
        """
        self.progress_callback = progress_callback
        self.log_callback = log_callback

        # Tracking data
        self.style_id_map = {}  # Maps old style ID -> new unique ID
        self.num_id_map = {}  # Maps old numbering ID -> new unique ID
        self.rel_id_map = {}  # Maps old relationship ID -> new unique ID

        # Counters for unique ID generation
        self.style_counter = 1000
        self.num_counter = 1000
        self.rel_counter = 1000

        # Statistics
        self.stats = {
            "total_files": 0,
            "merged_files": 0,
            "failed_files": 0,
            "total_styles": 0,
            "remapped_styles": 0,
            "total_numbering": 0,
            "remapped_numbering": 0,
            "errors": []
        }

    def _log(self, message: str, level: str = "INFO"):
        """Log a message"""
        if self.log_callback:
            self.log_callback(message, level)
        else:
            timestamp = time.strftime("%H:%M:%S")
            print(f"[{timestamp}] [{level}] {message}")

    def _update_progress(self, progress: float, message: str = ""):
        """Update progress"""
        if self.progress_callback:
            self.progress_callback(progress, message)

    def validate_docx(self, file_path: str) -> Tuple[bool, List[str]]:
        """
        Validate a DOCX file before merging.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Tuple of (is_valid, list_of_issues)
        """
        issues = []

        try:
            # Check if file exists
            if not os.path.exists(file_path):
                issues.append(f"File not found: {file_path}")
                return False, issues

            # Check if it's a valid ZIP file (DOCX is ZIP-based)
            if not zipfile.is_zipfile(file_path):
                issues.append(f"Not a valid DOCX file (invalid ZIP): {file_path}")
                return False, issues

            # Open as ZIP and check required files
            with zipfile.ZipFile(file_path, 'r') as zf:
                file_list = zf.namelist()

                # Check for required files
                required_files = ['[Content_Types].xml', 'word/document.xml']
                for req_file in required_files:
                    if req_file not in file_list:
                        issues.append(f"Missing required file: {req_file}")

                # Try to parse document.xml
                try:
                    doc_xml = zf.read('word/document.xml')
                    ET.fromstring(doc_xml)
                except ET.ParseError as e:
                    issues.append(f"Invalid XML in document.xml: {str(e)}")
                except Exception as e:
                    issues.append(f"Cannot read document.xml: {str(e)}")

                # Check styles.xml if exists
                if 'word/styles.xml' in file_list:
                    try:
                        styles_xml = zf.read('word/styles.xml')
                        ET.fromstring(styles_xml)
                    except ET.ParseError as e:
                        issues.append(f"Invalid XML in styles.xml: {str(e)}")

                # Check numbering.xml if exists
                if 'word/numbering.xml' in file_list:
                    try:
                        num_xml = zf.read('word/numbering.xml')
                        ET.fromstring(num_xml)
                    except ET.ParseError as e:
                        issues.append(f"Invalid XML in numbering.xml: {str(e)}")

            # Try to open with python-docx
            try:
                doc = Document(file_path)
                # Access basic properties to ensure it can be read
                _ = len(doc.paragraphs)
            except Exception as e:
                issues.append(f"Cannot open with python-docx: {str(e)}")

        except Exception as e:
            issues.append(f"Unexpected error during validation: {str(e)}")

        is_valid = len(issues) == 0
        return is_valid, issues

    def scan_and_validate_files(self, files: List[str]) -> Dict[str, any]:
        """
        Scan and validate all files before merging.

        Args:
            files: List of file paths

        Returns:
            Dictionary with validation results
        """
        self._log("=" * 60)
        self._log("SCANNING AND VALIDATING FILES")
        self._log("=" * 60)

        results = {
            "valid_files": [],
            "invalid_files": [],
            "total_issues": 0,
            "can_proceed": False
        }

        for i, file_path in enumerate(files, 1):
            self._log(f"[{i}/{len(files)}] Validating: {os.path.basename(file_path)}")
            self._update_progress((i - 0.5) / len(files) * 0.1, f"Validating file {i}/{len(files)}")

            is_valid, issues = self.validate_docx(file_path)

            if is_valid:
                self._log(f"  ✓ Valid", "SUCCESS")
                results["valid_files"].append(file_path)
            else:
                self._log(f"  ✗ Invalid - {len(issues)} issue(s)", "ERROR")
                for issue in issues:
                    self._log(f"    - {issue}", "ERROR")
                results["invalid_files"].append({
                    "file": file_path,
                    "issues": issues
                })
                results["total_issues"] += len(issues)

        # Can proceed if at least 2 files are valid
        results["can_proceed"] = len(results["valid_files"]) >= 2

        self._log("-" * 60)
        self._log(f"Validation complete: {len(results['valid_files'])} valid, "
                  f"{len(results['invalid_files'])} invalid")

        if results["can_proceed"]:
            self._log("✓ Ready to merge", "SUCCESS")
        else:
            self._log("✗ Cannot proceed - need at least 2 valid files", "ERROR")

        return results

    def _generate_unique_style_id(self, original_id: str, file_index: int) -> str:
        """Generate a unique style ID"""
        unique_id = f"Style{self.style_counter}_{file_index}"
        self.style_counter += 1
        return unique_id

    def _generate_unique_num_id(self, original_id: str, file_index: int) -> str:
        """Generate a unique numbering ID"""
        unique_id = f"{self.num_counter + file_index * 100}"
        self.num_counter += 1
        return unique_id

    def _generate_unique_rel_id(self, original_id: str, file_index: int) -> str:
        """Generate a unique relationship ID"""
        unique_id = f"rId{self.rel_counter}"
        self.rel_counter += 1
        return unique_id

    def _extract_styles(self, docx_path: str, file_index: int) -> Tuple[Optional[ET.Element], Dict[str, str]]:
        """
        Extract styles from a DOCX file and create ID mapping.

        Args:
            docx_path: Path to DOCX file
            file_index: Index of file in merge sequence

        Returns:
            Tuple of (styles_root_element, id_mapping)
        """
        id_mapping = {}

        with zipfile.ZipFile(docx_path, 'r') as zf:
            if 'word/styles.xml' not in zf.namelist():
                return None, id_mapping

            styles_xml = zf.read('word/styles.xml')
            root = ET.fromstring(styles_xml)

            # Find all style definitions
            for style in root.findall('.//w:style', NAMESPACES):
                style_id = style.get(f"{{{NAMESPACES['w']}}}styleId")
                if style_id:
                    # Generate unique ID
                    new_id = self._generate_unique_style_id(style_id, file_index)
                    id_mapping[style_id] = new_id

                    # Update the style ID in XML
                    style.set(f"{{{NAMESPACES['w']}}}styleId", new_id)

                    # Update basedOn references
                    based_on = style.find('.//w:basedOn', NAMESPACES)
                    if based_on is not None:
                        old_ref = based_on.get(f"{{{NAMESPACES['w']}}}val")
                        if old_ref in id_mapping:
                            based_on.set(f"{{{NAMESPACES['w']}}}val", id_mapping[old_ref])

                    # Update next style references
                    next_style = style.find('.//w:next', NAMESPACES)
                    if next_style is not None:
                        old_ref = next_style.get(f"{{{NAMESPACES['w']}}}val")
                        if old_ref in id_mapping:
                            next_style.set(f"{{{NAMESPACES['w']}}}val", id_mapping[old_ref])

            self.stats["total_styles"] += len(id_mapping)
            self.stats["remapped_styles"] += len(id_mapping)

        return root, id_mapping

    def _extract_numbering(self, docx_path: str, file_index: int) -> Tuple[Optional[ET.Element], Dict[str, str]]:
        """
        Extract numbering definitions and create ID mapping.

        Args:
            docx_path: Path to DOCX file
            file_index: Index of file in merge sequence

        Returns:
            Tuple of (numbering_root_element, id_mapping)
        """
        id_mapping = {}

        with zipfile.ZipFile(docx_path, 'r') as zf:
            if 'word/numbering.xml' not in zf.namelist():
                return None, id_mapping

            num_xml = zf.read('word/numbering.xml')
            root = ET.fromstring(num_xml)

            # Map abstract numbering IDs
            for abstract_num in root.findall('.//w:abstractNum', NAMESPACES):
                num_id = abstract_num.get(f"{{{NAMESPACES['w']}}}abstractNumId")
                if num_id:
                    new_id = self._generate_unique_num_id(num_id, file_index)
                    id_mapping[f"abstract_{num_id}"] = new_id
                    abstract_num.set(f"{{{NAMESPACES['w']}}}abstractNumId", new_id)

            # Map numbering instance IDs
            for num in root.findall('.//w:num', NAMESPACES):
                num_id = num.get(f"{{{NAMESPACES['w']}}}numId")
                if num_id:
                    new_id = self._generate_unique_num_id(num_id, file_index)
                    id_mapping[f"num_{num_id}"] = new_id
                    num.set(f"{{{NAMESPACES['w']}}}numId", new_id)

                    # Update reference to abstract numbering
                    abstract_ref = num.find('.//w:abstractNumId', NAMESPACES)
                    if abstract_ref is not None:
                        old_ref = abstract_ref.get(f"{{{NAMESPACES['w']}}}val")
                        if f"abstract_{old_ref}" in id_mapping:
                            abstract_ref.set(f"{{{NAMESPACES['w']}}}val",
                                           id_mapping[f"abstract_{old_ref}"])

            self.stats["total_numbering"] += len(id_mapping)
            self.stats["remapped_numbering"] += len(id_mapping)

        return root, id_mapping

    def _update_document_references(self, doc_xml: bytes, style_mapping: Dict[str, str],
                                     num_mapping: Dict[str, str]) -> bytes:
        """
        Update style and numbering references in document.xml.

        Args:
            doc_xml: Document XML as bytes
            style_mapping: Style ID mapping
            num_mapping: Numbering ID mapping

        Returns:
            Updated XML as bytes
        """
        root = ET.fromstring(doc_xml)

        # Update paragraph style references
        for para in root.findall('.//w:pStyle', NAMESPACES):
            old_val = para.get(f"{{{NAMESPACES['w']}}}val")
            if old_val in style_mapping:
                para.set(f"{{{NAMESPACES['w']}}}val", style_mapping[old_val])

        # Update character style references
        for run_style in root.findall('.//w:rStyle', NAMESPACES):
            old_val = run_style.get(f"{{{NAMESPACES['w']}}}val")
            if old_val in style_mapping:
                run_style.set(f"{{{NAMESPACES['w']}}}val", style_mapping[old_val])

        # Update table style references
        for tbl_style in root.findall('.//w:tblStyle', NAMESPACES):
            old_val = tbl_style.get(f"{{{NAMESPACES['w']}}}val")
            if old_val in style_mapping:
                tbl_style.set(f"{{{NAMESPACES['w']}}}val", style_mapping[old_val])

        # Update numbering references
        for num_pr in root.findall('.//w:numId', NAMESPACES):
            old_val = num_pr.get(f"{{{NAMESPACES['w']}}}val")
            if f"num_{old_val}" in num_mapping:
                num_pr.set(f"{{{NAMESPACES['w']}}}val", num_mapping[f"num_{old_val}"])

        return ET.tostring(root, encoding='utf-8', xml_declaration=True)

    def _copy_page_margins_to_result(self, source_file: str, result_file: str):
        """
        Copy page margin settings from source file to result file.
        This ensures the merged document has consistent page settings.

        Args:
            source_file: Path to source DOCX (first input file)
            result_file: Path to result DOCX file to update
        """
        try:
            from docx import Document

            self._log("Copying page margin settings to result file...", "INFO")

            # Load source and result documents
            source_doc = Document(source_file)
            result_doc = Document(result_file)

            if len(source_doc.sections) == 0 or len(result_doc.sections) == 0:
                self._log("  No sections found, skipping page margin copy", "WARNING")
                return

            source_section = source_doc.sections[0]

            # Apply to all sections in result
            for section in result_doc.sections:
                section.top_margin = source_section.top_margin
                section.bottom_margin = source_section.bottom_margin
                section.left_margin = source_section.left_margin
                section.right_margin = source_section.right_margin
                section.gutter = source_section.gutter

            # Save the updated result
            result_doc.save(result_file)
            self._log("  ✓ Page margins copied successfully", "SUCCESS")

        except Exception as e:
            self._log(f"  Warning: Could not copy page margins: {str(e)}", "WARNING")

    def _remap_styles_in_document(self, doc_path: str, file_index: int) -> Dict[str, str]:
        """
        Remap all style IDs in a document to prevent conflicts during merge.
        Creates a complete isolation of styles between documents.

        Args:
            doc_path: Path to the DOCX file
            file_index: Index of file in merge sequence

        Returns:
            Dictionary mapping old IDs to new IDs
        """
        import zipfile
        import tempfile
        import shutil

        style_mapping = {}

        try:
            # Extract and remap styles.xml
            styles_root, styles_map = self._extract_styles(doc_path, file_index)
            style_mapping.update(styles_map)

            # Extract and remap numbering.xml
            num_root, num_map = self._extract_numbering(doc_path, file_index)

            # Update document.xml with new IDs
            with zipfile.ZipFile(doc_path, 'r') as zf:
                doc_xml = zf.read('word/document.xml')
                updated_doc_xml = self._update_document_references(doc_xml, styles_map, num_map)

                # Create temporary copy with updated files
                temp_dir = tempfile.mkdtemp()
                temp_docx = os.path.join(temp_dir, 'temp.docx')

                with zipfile.ZipFile(temp_docx, 'w', zipfile.ZIP_DEFLATED) as new_zf:
                    # Copy all files
                    for item in zf.namelist():
                        if item == 'word/document.xml':
                            new_zf.writestr(item, updated_doc_xml)
                        elif item == 'word/styles.xml' and styles_root is not None:
                            new_zf.writestr(item, ET.tostring(styles_root, encoding='utf-8', xml_declaration=True))
                        elif item == 'word/numbering.xml' and num_root is not None:
                            new_zf.writestr(item, ET.tostring(num_root, encoding='utf-8', xml_declaration=True))
                        else:
                            new_zf.writestr(item, zf.read(item))

                # Replace original with updated version
                shutil.copy(temp_docx, doc_path)
                shutil.rmtree(temp_dir)

            self._log(f"  Remapped {len(styles_map)} styles for file #{file_index}", "INFO")

        except Exception as e:
            self._log(f"  Warning: Could not remap styles: {str(e)}", "WARNING")

        return style_mapping

    def merge_with_conflict_resolution(self, files: List[str], output_file: str) -> bool:
        """
        Merge DOCX files with automatic conflict resolution.

        Args:
            files: List of DOCX file paths to merge
            output_file: Output file path

        Returns:
            bool: True if successful, False otherwise
        """
        start_time = time.time()

        try:
            self._log("=" * 60)
            self._log(f"STARTING XML MERGE WITH CONFLICT RESOLUTION")
            self._log(f"Files to merge: {len(files)}")
            self._log("=" * 60)

            self.stats["total_files"] = len(files)

            # Step 1: Validate all files (10% progress)
            validation_result = self.scan_and_validate_files(files)
            if not validation_result["can_proceed"]:
                self._log("✗ Validation failed - cannot proceed", "ERROR")
                return False

            valid_files = validation_result["valid_files"]
            self._update_progress(0.1, "Validation complete")

            # Step 2: Use docxcompose with our enhanced handling (90% progress)
            self._log("-" * 60)
            self._log("MERGING DOCUMENTS")
            self._log("-" * 60)

            from docxcompose.composer import Composer

            # Load master document
            self._log(f"[1/{len(valid_files)}] Loading master: {os.path.basename(valid_files[0])}")
            self._update_progress(0.15, "Loading master document")

            t_start = time.time()
            master = Document(valid_files[0])
            composer = Composer(master)
            self.stats["merged_files"] = 1

            self._log(f"  ✓ Loaded in {time.time() - t_start:.2f}s", "SUCCESS")

            # Merge remaining files with style isolation
            progress_per_file = 0.7 / (len(valid_files) - 1)
            current_progress = 0.15

            for i, file_path in enumerate(valid_files[1:], 2):
                try:
                    self._log(f"[{i}/{len(valid_files)}] Merging: {os.path.basename(file_path)}")
                    t_file_start = time.time()

                    # Remap styles to prevent conflicts
                    self._log(f"  Remapping styles for isolation...", "INFO")
                    style_map = self._remap_styles_in_document(file_path, i)

                    # Load document
                    doc = Document(file_path)

                    # Append to master
                    composer.append(doc)
                    self.stats["merged_files"] += 1

                    elapsed = time.time() - t_file_start
                    self._log(f"  ✓ Merged in {elapsed:.2f}s", "SUCCESS")

                    current_progress += progress_per_file
                    self._update_progress(current_progress, f"Merged {i}/{len(valid_files)} files")

                except Exception as e:
                    self._log(f"  ✗ Failed: {str(e)}", "ERROR")
                    self.stats["failed_files"] += 1
                    self.stats["errors"].append(f"{os.path.basename(file_path)}: {str(e)}")

            # Step 3: Save result (10% progress)
            self._log("-" * 60)
            self._log("Saving merged document...")
            self._update_progress(0.9, "Saving result")

            t_save = time.time()
            composer.save(output_file)
            save_time = time.time() - t_save

            self._log(f"✓ Saved in {save_time:.2f}s", "SUCCESS")

            # Step 4: Copy page margins from first file to result
            self._log("-" * 60)
            self._copy_page_margins_to_result(valid_files[0], output_file)

            self._update_progress(1.0, "Complete")

            # Summary
            total_time = time.time() - start_time
            self._log("=" * 60)
            self._log("MERGE COMPLETE", "SUCCESS")
            self._log("=" * 60)
            self._log(f"Total files: {self.stats['total_files']}")
            self._log(f"Merged successfully: {self.stats['merged_files']}")
            self._log(f"Failed: {self.stats['failed_files']}")
            self._log(f"Total time: {total_time:.2f}s")
            self._log(f"Output: {output_file}")
            self._log(f"Output size: {os.path.getsize(output_file) / 1024:.1f} KB")

            if self.stats["errors"]:
                self._log("-" * 60)
                self._log("ERRORS:")
                for error in self.stats["errors"]:
                    self._log(f"  - {error}", "ERROR")

            return True

        except Exception as e:
            self._log(f"✗ CRITICAL ERROR: {str(e)}", "ERROR")
            import traceback
            traceback.print_exc()
            self.stats["errors"].append(f"Critical: {str(e)}")
            return False


def merge_docx_advanced(files: List[str], output_file: str,
                        progress_callback: Optional[Callable] = None,
                        log_callback: Optional[Callable] = None) -> bool:
    """
    Convenience function for merging DOCX files with advanced features.

    Args:
        files: List of DOCX file paths to merge
        output_file: Output file path
        progress_callback: Optional callback for progress updates
        log_callback: Optional callback for log messages

    Returns:
        bool: True if successful, False otherwise
    """
    merger = XMLAdvancedMerger(progress_callback, log_callback)
    return merger.merge_with_conflict_resolution(files, output_file)


if __name__ == "__main__":
    import sys

    # Test with command line arguments or default files
    if len(sys.argv) > 2:
        input_files = sys.argv[1:-1]
        output_filename = sys.argv[-1]
    else:
        # Default test
        input_files = ["1.docx", "2.docx", "3.docx", "4.docx"]
        output_filename = "merged_advanced.docx"

    print(f"Merging {len(input_files)} files...")
    print(f"Output: {output_filename}\n")

    success = merge_docx_advanced(input_files, output_filename)

    if success:
        print(f"\n✅ Success! Check: {output_filename}")
    else:
        print(f"\n❌ Failed!")
        sys.exit(1)

