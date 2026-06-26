# -*- coding: utf-8 -*-
"""
xml_merger.py
-------------
XML-based document merger for combining multiple Word documents.

Optimized with progress tracking and logging support for multi-threaded execution.
Uses docxcompose with fixes for section handling.

Supports two merge modes:
1. Standard mode: Fast merge using docxcompose (default)
2. Advanced mode: Merge with validation and conflict resolution
"""

from docx import Document
from docxcompose.composer import Composer
import docxcompose.composer as composer_module
import time
from typing import List, Optional, Callable
from pathlib import Path
import os


# Monkey patch to fix section type issues
def _noop_fix_section_types(self, doc):
    """
    No-op replacement for fix_section_types to avoid IndexError.
    This prevents docxcompose from incorrectly handling section types.
    """
    return

composer_module.Composer.fix_section_types = _noop_fix_section_types


class XMLDocumentMerger:
    """
    Handles merge of multiple Word documents using docxcompose.
    Supports progress callbacks and detailed logging.

    Can use standard or advanced merge mode:
    - Standard: Fast merge (default)
    - Advanced: With validation and conflict resolution
    """

    def __init__(self, progress_callback: Optional[Callable] = None,
                 log_callback: Optional[Callable] = None,
                 use_advanced: bool = False):
        """
        Initialize the XML document merger.

        Args:
            progress_callback: Callback function(progress: float, message: str)
            log_callback: Callback function(message: str, level: str)
            use_advanced: Use advanced merger with conflict resolution
        """
        self.progress_callback = progress_callback
        self.log_callback = log_callback
        self.use_advanced = use_advanced

        # Statistics
        self.stats = {
            "total_files": 0,
            "processed_files": 0,
            "failed_files": 0,
            "start_time": None,
            "end_time": None
        }

    def log(self, message: str, level: str = "INFO"):
        """Log a message if callback is provided"""
        if self.log_callback:
            self.log_callback(message, level)
        else:
            print(f"[{level}] {message}")

    def update_progress(self, progress: float, message: str = ""):
        """Update progress if callback is provided"""
        if self.progress_callback:
            self.progress_callback(progress, message)

    def validate_files(self, file_list: List[str]) -> tuple:
        """
        Validate that all files exist and are accessible.

        Args:
            file_list: List of file paths to validate

        Returns:
            Tuple of (valid_files, invalid_files)
        """
        valid = []
        invalid = []

        for file_path in file_list:
            if os.path.exists(file_path):
                try:
                    # Try to open briefly to check if it's a valid docx
                    doc = Document(file_path)
                    valid.append(file_path)
                    self.log(f"Validated: {Path(file_path).name}", "DEBUG")
                except Exception as e:
                    invalid.append((file_path, str(e)))
                    self.log(f"Invalid file {file_path}: {e}", "WARNING")
            else:
                invalid.append((file_path, "File not found"))
                self.log(f"File not found: {file_path}", "WARNING")

        return valid, invalid

    def merge_documents(self, file_list: List[str], output_path: str,
                       validate_first: bool = True) -> dict:
        """
        Merge multiple Word documents into one.

        Args:
            file_list: List of file paths to merge
            output_path: Output file path
            validate_first: Whether to validate files before merge

        Returns:
            Dictionary with results:
                - success: bool
                - output_file: str
                - files_merged: int
                - files_failed: int
                - elapsed_time: float
                - error: str (if failed)
        """
        # Use advanced merger if requested
        if self.use_advanced:
            try:
                from .xml_advanced_merger import XMLAdvancedMerger

                self.log("Using advanced merger with conflict resolution", "INFO")

                advanced_merger = XMLAdvancedMerger(
                    progress_callback=self.progress_callback,
                    log_callback=self.log_callback
                )

                success = advanced_merger.merge_with_conflict_resolution(file_list, output_path)

                # Convert advanced merger stats to standard format
                return {
                    "success": success,
                    "output_file": output_path,
                    "files_merged": advanced_merger.stats.get("merged_files", 0),
                    "files_failed": advanced_merger.stats.get("failed_files", 0),
                    "elapsed_time": 0,  # Advanced merger tracks this internally
                }
            except ImportError:
                self.log("Advanced merger not available, falling back to standard mode", "WARNING")

        # Standard merge mode
        self.stats["start_time"] = time.time()
        self.stats["total_files"] = len(file_list)

        try:
            self.log("=" * 60, "INFO")
            self.log(f"Starting merge process for {len(file_list)} files", "INFO")
            self.log("=" * 60, "INFO")

            self.update_progress(0, "Initializing merge process")

            # Validate input
            if not file_list or len(file_list) == 0:
                error_msg = "File list is empty"
                self.log(error_msg, "ERROR")
                return {"success": False, "error": error_msg}

            # Validate files if requested
            if validate_first:
                self.log("Validating input files...", "INFO")
                self.update_progress(5, "Validating files")

                valid_files, invalid_files = self.validate_files(file_list)

                if invalid_files:
                    self.log(f"Found {len(invalid_files)} invalid files:", "WARNING")
                    for file_path, reason in invalid_files:
                        self.log(f"  - {file_path}: {reason}", "WARNING")

                if not valid_files:
                    error_msg = "No valid files to merge"
                    self.log(error_msg, "ERROR")
                    return {"success": False, "error": error_msg}

                if len(valid_files) < len(file_list):
                    self.log(f"Proceeding with {len(valid_files)} valid files", "WARNING")
                    file_list = valid_files

            # Load master document
            self.log(f"Loading master document: {Path(file_list[0]).name}", "INFO")
            self.update_progress(10, f"Loading master: {Path(file_list[0]).name}")

            load_start = time.time()
            master = Document(file_list[0])
            composer = Composer(master)
            load_time = time.time() - load_start

            self.log(f"Master loaded in {load_time:.4f}s", "SUCCESS")
            self.stats["processed_files"] = 1

            # Calculate progress increments
            if len(file_list) > 1:
                progress_per_file = 70.0 / (len(file_list) - 1)
            else:
                progress_per_file = 0

            # Merge remaining files
            for i, file_path in enumerate(file_list[1:], 1):
                file_start = time.time()
                file_name = Path(file_path).name

                current_progress = 10 + (i * progress_per_file)
                self.update_progress(
                    current_progress,
                    f"Merging file {i}/{len(file_list)-1}: {file_name}"
                )

                self.log(f"[{i}/{len(file_list)-1}] Processing: {file_name}", "INFO")

                try:
                    # Load document
                    doc = Document(file_path)

                    # Append to composer
                    composer.append(doc)

                    file_time = time.time() - file_start
                    self.log(f"  ✓ Merged in {file_time:.4f}s", "SUCCESS")
                    self.stats["processed_files"] += 1

                except Exception as e:
                    self.log(f"  ✗ Failed to merge {file_name}: {e}", "ERROR")
                    self.stats["failed_files"] += 1

            # Save merged document
            self.log("Saving merged document (this may take a while)...", "INFO")
            self.update_progress(85, "Saving merged document")

            save_start = time.time()

            # Ensure output directory exists
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)

            composer.save(output_path)
            save_time = time.time() - save_start

            self.log(f"Document saved in {save_time:.4f}s", "SUCCESS")

            # Final statistics
            self.stats["end_time"] = time.time()
            elapsed_time = self.stats["end_time"] - self.stats["start_time"]

            self.update_progress(100, "Merge completed")

            self.log("=" * 60, "INFO")
            self.log("MERGE COMPLETED SUCCESSFULLY", "SUCCESS")
            self.log(f"Total files: {self.stats['total_files']}", "INFO")
            self.log(f"Successfully merged: {self.stats['processed_files']}", "INFO")
            self.log(f"Failed: {self.stats['failed_files']}", "INFO")
            self.log(f"Total time: {elapsed_time:.4f} seconds", "INFO")
            self.log(f"Output file: {output_path}", "INFO")
            self.log("=" * 60, "INFO")

            return {
                "success": True,
                "output_file": output_path,
                "files_merged": self.stats["processed_files"],
                "files_failed": self.stats["failed_files"],
                "elapsed_time": elapsed_time,
                "save_time": save_time
            }

        except Exception as e:
            self.stats["end_time"] = time.time()
            error_msg = f"Critical error during merge: {e}"
            self.log(error_msg, "ERROR")

            return {
                "success": False,
                "error": error_msg,
                "files_merged": self.stats["processed_files"],
                "files_failed": self.stats["failed_files"]
            }


# Worker function for threading integration
def xml_merge_worker(thread, file_list: List[str], output_file: str,
                    validate_files: bool = True) -> dict:
    """
    Worker function for merge XML documents with thread support.

    Args:
        thread: WorkerThread instance
        file_list: List of file paths to merge
        output_file: Output file path
        validate_files: Whether to validate files first

    Returns:
        Dictionary with results
    """
    def progress_callback(progress: float, message: str):
        thread.update_progress(progress, message)

    def log_callback(message: str, level: str):
        thread.log(message, level)

    merger = XMLDocumentMerger(
        progress_callback=progress_callback,
        log_callback=log_callback
    )

    return merger.merge_documents(file_list, output_file, validate_files)

